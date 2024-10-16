import xml.etree.ElementTree as ET
import csv
import argparse
import os
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

def parse_nessus_file(nessus_file, port_filter, severity_filter, plugin_filter, output_file, sort_by, docx_output):
    tree = ET.parse(nessus_file)
    root = tree.getroot()

    result_list = []

    # Iterate through each ReportHost block
    for report_host in root.findall(".//ReportHost"):
        host_ip = None
        host_fqdn = None

        # Retrieve IP and FQDN from HostProperties block
        for tag in report_host.findall(".//HostProperties/tag"):
            if tag.get('name') == 'host-ip':
                host_ip = tag.text
            if tag.get('name') == 'host-fqdn':
                host_fqdn = tag.text

        # Iterate through ReportItems to check for ports and filter by severity or plugin
        for report_item in report_host.findall(".//ReportItem"):
            port = report_item.get('port')
            severity = int(report_item.get('severity'))  # Severity as an integer
            plugin_name = report_item.get('pluginName')

            # Check if the severity meets the filter and if the plugin name matches (if provided)
            if severity >= severity_filter and (not plugin_filter or plugin_filter.lower() in plugin_name.lower()):
                # Only filter by port if one is specified
                if not port_filter or port == port_filter:
                    findings_text = plugin_name
                    fqdn_text = host_fqdn if host_fqdn else ""
                    result_list.append([host_ip, fqdn_text, port, findings_text])

    # Remove exact duplicate rows
    result_list = list(map(list, set(map(tuple, result_list))))

    # Sort results if sort_by argument is provided
    if sort_by:
        sort_by = sort_by.lower()
        if sort_by == "ip":
            result_list.sort(key=lambda x: x[0])
        elif sort_by == "fqdn":
            result_list.sort(key=lambda x: x[1])
        elif sort_by == "port":
            result_list.sort(key=lambda x: int(x[2]))  # Ensure numeric sorting for ports
        elif sort_by == "findings":
            result_list.sort(key=lambda x: x[3])

    # Determine file format based on extension (default to csv)
    if output_file and not docx_output:
        if output_file.endswith(".txt"):
            file_format = "txt"
        elif output_file.endswith(".csv") or not os.path.splitext(output_file)[1]:
            file_format = "csv"
        else:
            file_format = "csv"

        # Save the results to a file in the specified format
        if file_format == 'csv':
            with open(output_file, mode='w', newline='') as file:
                writer = csv.writer(file)
                writer.writerow(["IP", "FQDN", "Port", "Findings"])
                writer.writerows(result_list)
        else:  # Save as txt
            with open(output_file, mode='w') as file:
                for result in result_list:
                    file.write(f"{result[0]},{result[1]},{result[2]},{result[3]}\n")

        print(f"Results saved to {output_file}")

    # Handle .docx output
    if docx_output:
        document = Document()

        # Create a title for the document
        document.add_heading('Nessus Results', 0)

        # Create the table with headers
        table = document.add_table(rows=1, cols=4)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'IP'
        hdr_cells[1].text = 'FQDN'
        hdr_cells[2].text = 'Port'
        hdr_cells[3].text = 'Findings'

        # Apply basic formatting to the header row
        for cell in hdr_cells:
            shading_elm = parse_xml(r'<w:shd {} w:fill="B7DEE8"/>'.format(nsdecls('w')))
            cell._element.get_or_add_tcPr().append(shading_elm)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(10)

        # Add data rows
        for result in result_list:
            row_cells = table.add_row().cells
            row_cells[0].text = result[0]
            row_cells[1].text = result[1]
            row_cells[2].text = result[2]
            row_cells[3].text = result[3]

        # Save document
        docx_file = output_file if output_file.endswith(".docx") else output_file + ".docx"
        document.save(docx_file)
        print(f"Results saved to {docx_file}")

# Argument parsing for command-line usage
def main():
    parser = argparse.ArgumentParser(description="Parse Nessus XML file to extract data by port, severity, plugin name, and export to CSV, TXT, or DOCX.")
    parser.add_argument('-i', '--input', required=True, help="Path to the Nessus XML file")
    parser.add_argument('-p', '--port', help="Port to filter (e.g., '80'). If not set, processes all ports.")
    parser.add_argument('-s', '--severity', type=int, default=0, help="Severity filter (0=Info, 1=Low, 2=Medium, 3=High, 4=Critical)")
    parser.add_argument('-n', '--plugin', help="Plugin name filter (optional)")
    parser.add_argument('-o', '--output', required=True, help="Output file path (.txt, .csv, or .docx)")
    parser.add_argument('--sort-by', help="Column to sort by: 'ip', 'fqdn', 'port', or 'findings' (optional)")
    parser.add_argument('--docx', action='store_true', help="Generate DOCX output (optional)")

    args = parser.parse_args()

    # Run the parser function with provided arguments
    parse_nessus_file(args.input, args.port, args.severity, args.plugin, args.output, args.sort_by, args.docx)

if __name__ == "__main__":
    main()
